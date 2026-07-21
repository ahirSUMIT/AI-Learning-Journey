import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("Where is the API KEY")

client=Groq(api_key=my_api_key)
model = "llama-3.3-70b-versatile"
REQUEST_DELAY = 2


job_description="""
Description
Do you want to solve real customer problems through innovative technology? Do you enjoy working on scalable services in a collaborative team environment? Do you want to see your code directly impact millions of customers worldwide?

At Amazon, we hire the best minds in technology to innovate and build on behalf of our customers. Customer obsession is part of our company DNA, which has made us one of the world's most beloved brands.

Our Software Development Engineers (SDEs) use modern technology to solve complex problems while seeing their work's impact first-hand. The challenges SDEs solve at Amazon are meaningful and influence millions of customers, sellers, and products globally. We seek individuals passionate about creating new products, features, and services while managing ambiguity in an environment where development cycles are measured in weeks, not years.

At Amazon, we believe in ownership at every level. As an SDE-I, you'll own the entire lifecycle of your code - from design through deployment and ongoing operations. This ownership mindset, combined with our commitment to operational excellence, ensures we deliver the highest quality solutions for our customers.

We're looking for curious minds who think big and want to define tomorrow's technology. At Amazon, you'll grow into the high-impact engineer you know you can be, supported by a culture of learning and mentorship. Every day brings exciting new challenges and opportunities for personal growth.
Key job responsibilities
• Collaborate and communicate effectively with experienced cross-disciplinary Amazonians to design, build, and operate innovative products and services that delight our customers, while participating in technical discussions to drive solutions forward.
• Design and develop scalable solutions using cloud-native architectures and microservices in a large distributed computing environment.
• Participate in code reviews and contribute to technical documentation.
• Build and maintain resilient distributed systems that are scalable, fault-tolerant, and cost-effective.
• Leverage and contribute to the development of GenAI and AI-powered tools to enhance development productivity while staying current with emerging technologies.
• Write clean, maintainable code following best practices and design patterns.
• Work in an agile environment practicing CI/CD principles while participating in operational responsibilities including on-call duties.
• Demonstrate operational excellence through monitoring, troubleshooting, and resolving production issues.
Basic Qualifications
- Experience with at least one general-purpose programming language such as Java, Python, C++, C#, Go, Rust, or TypeScript
- Experience with data structure implementation, basic algorithm development, and/or object-oriented design principles
- Currently has, or is in the process of obtaining a bachelor’s degree in Computer Science, Computer Engineering, Data Science, Information Systems, or related STEM fields
- Must be 18 years of age of older
Preferred Qualifications
- Experience from previous technical internship(s) or demonstrated project experience
- Experience with one or more of the following: AI tools for development productivity, Cloud platforms (preferably AWS), Database systems (SQL and NoSQL), Contributing to open-source projects, Version control systems, Debugging and troubleshooting complex systems
- Demonstrated ability to learn and adapt to new technologies quickly
- Basic understanding of software development lifecycle (SDLC)
- Strong problem-solving and analytical skills
- Excellent written and verbal communication skills
"""
class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""
message_system={
    "role" : "system",
    "content" : system_prompt
}
message_user={
    "role" : "user",
    "content" : user_prompt
}
response_format={
    "type" : "json_object"
}


messages=[message_system, message_user]

response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)


answer=response.choices[0].message.content

raw_json=answer
# print(raw_json)



import json
job_data=json.loads(raw_json)

job = JobD(**job_data)

print(job.minimum_experience)
print(job.education_requirements)



#parse real
class MatchDetails(BaseModel):
    candidate_name: str | None = None
    matching_skills: list[str]
    missing_important_skills: list[str]
    experience_requirement_met: bool | None = None
    verdict: str

class MatchResult(BaseModel):
    score: int = Field(
        ge=0,
        le=100,
        description="Overall match score from 0 to 100"
    )
    details: MatchDetails
    
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = Field(default_factory=list)  #to define mutable defaults.

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an ATS HR recruiter.

    Compare the following resume with the job description.

    JOB DESCRIPTION:

    {job.model_dump_json()}

    RESUME:

    {resume.model_dump_json()}

    Return exactly ONE valid JSON object matching this schema:

    {match_schema}

    Scoring Rules:

    1. Score must be an integer between 0 and 100.

    2. Score ONLY against the requirements present in the job description.

    3. Do NOT give extra marks for unrelated skills.

    4. Do NOT infer skills that are not explicitly mentioned in the resume.

    5. matching_skills must contain ONLY the ACTUAL skill names found in the resume.

    Examples:
    ✓ Java
    ✓ Python
    ✓ C++
    ✓ Git
    ✓ SQL
    ✓ AWS

    Never return requirement sentences such as:
    ✗ Experience with at least one general-purpose programming language
    ✗ Experience with data structure implementation

    6. missing_important_skills must contain ONLY required skills from the job description that are absent from the resume.

    7. Ignore unrelated technologies and certifications.

    8. Use exactly the field names defined in the schema.

    9. Never rename keys.

    10. Never create extra keys.

    11. Never omit required keys.

    12. If information is unavailable:
        - string -> ""
        - list -> []
        - bool -> false

    Return ONLY the JSON object.
    
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format, temperature=0)
    print(response.choices[0].message.content)
    data = json.loads(response.choices[0].message.content)
    details = data["details"]

    if "name" in details:
        details["candidate_name"] = details.pop("name")

    if "candidate" in details:
        details["candidate_name"] = details.pop("candidate")

    if "final_verdict" in details:
        details["verdict"] = details.pop("final_verdict")

    if "experience_met" in details:
        details["experience_requirement_met"] = details.pop("experience_met")

    if "missing_skills" in details:
        details["missing_important_skills"] = details.pop("missing_skills")
    
    data["details"] = details
    # Normalize score
    if isinstance(data.get("score"), float) and data["score"] <= 1:
        data["score"] = round(data["score"] * 100)

    return MatchResult(**data)

def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    for exp in data.get("experiences", []):
        if exp.get("skills_used") is None:
            exp["skills_used"] = []
    resume = Resume(**data)
    return resume


from docx import Document
import fitz

def read_pdf(file_path):
    try:
        doc = fitz.open(file_path)

        text = ""

        for page in doc:
            text += page.get_text()

        doc.close()

        return text.strip()

    except Exception as e:
        print(f"❌ Error reading {file_path.name}: {e}")
        return None

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None



# lets do it now
resume_folder = Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():

    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue

    print("\nProcessing:", file_path.name)

    resume_text = read_resume(file_path)

    # Debug: Check extracted text
    print("=" * 50)
    print(f"Characters extracted: {len(resume_text) if resume_text else 0}")

    if resume_text:
        print(resume_text[:300])

    print("=" * 50)

    if not resume_text or len(resume_text.strip()) < 50:
        print(f"Skipping {file_path.name} (No readable text)")
        continue

    try:
        # First LLM call
        parsed_resume = parse_resume(resume_text)
        time.sleep(REQUEST_DELAY)

        # Second LLM call
        result = final_score(job, parsed_resume)
        time.sleep(REQUEST_DELAY)

    except Exception as e:
        print(f"❌ LLM failed for {file_path.name}: {e}")
        continue
    #score and details
    #account chatgpt
    #started sending request to millions
    #chattgpt server will jam 

    
    print("Score:", result.score)
    all_results.append({
    "name": parsed_resume.name,
    "score": result.score,
    "matching_skills": result.details.matching_skills,
    "missing_skills": result.details.missing_important_skills,
    "verdict": result.details.verdict
    })
if not all_results:
    print("❌ No resumes were successfully processed.")
    exit()
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
    )

top_2 = all_results[:2]

if len(all_results) > 2:
    worst_2 = all_results[-2:]
else:
    worst_2 = []


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print("Matching Skills :", candidate["matching_skills"])
    print("Missing Skills  :", candidate["missing_skills"])
    print("Verdict         :", candidate["verdict"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print("Matching Skills :", candidate["matching_skills"])
    print("Missing Skills  :", candidate["missing_skills"])
    print("Verdict         :", candidate["verdict"])
